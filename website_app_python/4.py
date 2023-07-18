import os
import xlsxwriter
import docx
from docx import Document
document = Document
with open("/Users/zhoujie/Desktop/data/ottolenghi 2/roast_chicken.txt") as p:
    text1= p.read()
    #print(text1)

text_list = text1.split("\n")
print(text_list.index("OTTOLENGHI’S ROAST CHICKEN WITH ZA’ATAR AND SUMAC"))
print(text_list.index("How you make it"))
print
context1 = text_list[1:5]
print(context1)
context2 = text_list[6:]
print(context2)
document = Document()
p = document.add_heading("OTTOLENGHI’S ROAST CHICKEN WITH ZA’ATAR AND SUMAC", 1)
p = document.add_heading("")

for x in context1:
    p = document.add_paragraph(x)
    p.add_run(x).italic = True
p = document.add_heading("How you make it", 2)
#string_context2 = "".join(context2)
word_list = ["onions", "lemon", "garlic", "bowl",
             "crisp", "garnish"]
word = ["chicken"]
number = 1
for y in context2:
    y = str(number) + ". " + y
    number += 1
    z = [i[0]+"_____" if i in word_list else i for i in y.split()]
    print(z)
    y = " ".join(z)
    p = document.add_paragraph(y)






"""
list_context2 = string_context2.split()
list_final = [i[0]+"____" if i in word_list else i for i in list_context2]
context2 = " ".join(list_final)
print(context2)
paragraph_list = context2.split(".")
number = 1
for y in paragraph_list:
    y = str(number) + ". " + y
    number += 1
    p = document.add_paragraph(y)
"""

document.save(os.path.join('/Users/zhoujie/Desktop/data/ottolenghi 2/roast_chicken.docx'))